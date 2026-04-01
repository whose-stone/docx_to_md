import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("Run: pip install python-docx")
    sys.exit(1)


def runs_to_markdown(paragraph):
    md = ""
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        md += text
    return md


def table_to_markdown(table):
    rows = []
    for row in table.rows:
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    col_count = max(len(r) for r in rows)
    for row in rows:
        while len(row) < col_count:
            row.append("")
    lines = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * col_count) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def get_list_info(paragraph):
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return False, False, 0
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return False, False, 0
    num_id_el = num_pr.find(qn("w:numId"))
    ilvl_el = num_pr.find(qn("w:ilvl"))
    if num_id_el is None:
        return False, False, 0
    num_id = int(num_id_el.get(qn("w:val"), 0))
    ilvl = int(ilvl_el.get(qn("w:val"), 0)) if ilvl_el is not None else 0
    if num_id == 0:
        return False, False, 0
    style_name = paragraph.style.name if paragraph.style else ""
    is_ordered = "List Number" in style_name
    return True, is_ordered, ilvl


AREAS_PATTERN = re.compile(r"(Areas\s+of\s+Law|Areas|Tags)\s*[:\-\u2013\u2014]\s*(.+)", re.IGNORECASE)


def extract_areas_of_law(text):
    match = AREAS_PATTERN.search(text)
    if not match:
        return []
    return [tag.strip() for tag in re.split(r"[,;]", match.group(2)) if tag.strip()]


def sanitize_tag(tag):
    return re.sub(r"\s+", "-", tag.strip())


def docx_to_markdown(docx_path):
    doc = Document(docx_path)
    lines = []
    all_tags = []
    ordered_counters = {}
    table_map = {table._tbl: table for table in doc.tables}
    para_map = {paragraph._p: paragraph for paragraph in doc.paragraphs}

    for child in doc.element.body:
        tag = child.tag.split("}")[-1]
        if tag == "tbl" and child in table_map:
            lines.append("")
            lines.append(table_to_markdown(table_map[child]))
            lines.append("")
            ordered_counters.clear()
            continue
        if tag == "p" and child in para_map:
            para = para_map[child]
            raw_text = para.text.strip()
            found_tags = extract_areas_of_law(raw_text)
            if found_tags:
                all_tags.extend(found_tags)
                lines.append("")
                lines.append(f"> **Areas of Law:** {', '.join(found_tags)}")
                lines.append("")
                continue

            style_name = para.style.name if para.style else "Normal"
            heading_match = re.match(r"Heading (\d+)", style_name)
            if heading_match:
                level = int(heading_match.group(1))
                md_text = runs_to_markdown(para)
                if md_text:
                    lines.append("")
                    lines.append("#" * level + " " + md_text)
                    lines.append("")
                ordered_counters.clear()
                continue

            if style_name in ("Title", "Subtitle"):
                md_text = runs_to_markdown(para)
                if md_text:
                    prefix = "# " if style_name == "Title" else "## "
                    lines.append("")
                    lines.append(prefix + md_text)
                    lines.append("")
                continue

            is_list, is_ordered, ilvl = get_list_info(para)
            if is_list:
                md_text = runs_to_markdown(para)
                indent = "  " * ilvl
                if is_ordered:
                    ordered_counters[ilvl] = ordered_counters.get(ilvl, 0) + 1
                    for depth in list(ordered_counters.keys()):
                        if depth > ilvl:
                            del ordered_counters[depth]
                    bullet = f"{ordered_counters[ilvl]}."
                else:
                    bullet = "-"
                    ordered_counters.clear()
                lines.append(f"{indent}{bullet} {md_text}")
                continue

            if not raw_text:
                if lines and lines[-1] != "":
                    lines.append("")
                ordered_counters.clear()
                continue

            md_text = runs_to_markdown(para)
            if md_text:
                lines.append(md_text)
            ordered_counters.clear()

    body_md = "\n".join(lines)
    body_md = re.sub(r"\n{3,}", "\n\n", body_md).strip()
    return body_md, all_tags


def build_frontmatter(tags, source_file):
    if not tags:
        yaml_tags = "  []"
    else:
        yaml_tags = "\n" + "\n".join(f"  - {sanitize_tag(tag)}" for tag in tags)
    return f"---\nsource: \"{source_file}\"\ntags:{yaml_tags}\n---\n"


def convert_docx_file(docx_path, source_name=None):
    docx_path = Path(docx_path)
    body, tags = docx_to_markdown(docx_path)
    source_name = source_name or docx_path.name
    markdown = build_frontmatter(tags, source_name) + "\n" + body + "\n"
    return {
        "source": docx_path,
        "source_name": source_name,
        "tags": tags,
        "markdown": markdown,
    }


def convert_all(source_root, logger=None):
    source_root = Path(source_root).resolve()
    output_root = source_root / "Conversions"
    output_root.mkdir(exist_ok=True)
    docx_files = list(source_root.rglob("*.docx"))

    def emit(message):
        if logger:
            logger(message)
        else:
            print(message)

    if not docx_files:
        emit(f"No .docx files found under: {source_root}")
        return {
            "source": source_root,
            "output": output_root,
            "found": 0,
            "converted": 0,
            "failed": 0,
            "results": [],
        }

    emit(f"Found {len(docx_files)} .docx file(s). Converting...")
    emit("")
    success = 0
    errors = 0
    results = []

    for docx_path in docx_files:
        if output_root in docx_path.parents:
            continue

        relative = docx_path.relative_to(source_root)
        md_path = output_root / relative.with_suffix(".md")
        md_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            converted = convert_docx_file(docx_path)
            md_path.write_text(converted["markdown"], encoding="utf-8")
            emit(f"  OK   {relative}")
            emit(f"       tags: {converted['tags'] if converted['tags'] else 'none found'}")
            success += 1
            results.append(
                {
                    "status": "ok",
                    "source": docx_path,
                    "output": md_path,
                    "relative": relative,
                    "tags": converted["tags"],
                }
            )
        except Exception as exc:
            emit(f"  FAIL {relative}  -->  {exc}")
            errors += 1
            results.append(
                {
                    "status": "fail",
                    "source": docx_path,
                    "output": md_path,
                    "relative": relative,
                    "error": str(exc),
                }
            )

    emit("")
    emit(f"Done. {success} converted, {errors} failed.")
    emit(f"Output: {output_root}")
    return {
        "source": source_root,
        "output": output_root,
        "found": len(docx_files),
        "converted": success,
        "failed": errors,
        "results": results,
    }


def run_cli(source):
    source = Path(source).resolve()
    if not source.is_dir():
        print(f"ERROR: '{source}' is not a valid directory.")
        return 1

    print(f"Source: {source}")
    print(f"Output: {source / 'Conversions'}")
    print("-" * 50)
    convert_all(source)
    return 0


if __name__ == "__main__":
    source = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else Path.cwd()
    sys.exit(run_cli(source))
